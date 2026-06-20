"""Export examination papers to Word (.docx) — board table layout."""

import re

ROMAN = ["i", "ii", "iii", "iv", "v", "vi", "vii", "viii", "ix", "x", "xi", "xii"]
FONT_NAME = "Times New Roman"


def _clean_mcq_option(opt, index: int) -> str:
    s = str(opt or "").strip()
    s = re.sub(r"^\(?[A-Da-d]\)?[.):\-\s]+", "", s)
    s = re.sub(r"^[A-Da-d]\s+", "", s)
    if re.fullmatch(r"[A-Da-d]", s):
        return f"Option {index + 1}"
    return s or f"Option {index + 1}"


def _section_key(sec):
    qtype = (sec.get("questionType") or "").lower()
    title = (sec.get("title") or "").lower()
    if qtype == "short" or "short" in title:
        return "short"
    if qtype == "long" or "long" in title:
        return "long"
    return "objective"


def _setup_document(doc):
    from docx.shared import Inches, Pt

    sec = doc.sections[0]
    sec.top_margin = Inches(0.55)
    sec.bottom_margin = Inches(0.55)
    sec.left_margin = Inches(0.65)
    sec.right_margin = Inches(0.65)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.15


def _shade_cell(cell, fill="E6E6E6"):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def _set_col_widths(table, widths_inches):
    from docx.shared import Inches

    for row in table.rows:
        for i, w in enumerate(widths_inches):
            if i < len(row.cells):
                row.cells[i].width = Inches(w)


def _para(cell_or_doc, text, *, bold=False, size=12, align=None, space_after=4):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    if hasattr(cell_or_doc, "add_paragraph"):
        p = cell_or_doc.add_paragraph()
    else:
        p = cell_or_doc.paragraphs[0]
        p.clear()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    return p


def _merge_row(table, row_idx, colspan=2):
    row = table.rows[row_idx]
    row.cells[0].merge(row.cells[colspan - 1])
    return row.cells[0]


def _add_mcq_options_cell(cell, options):
    from docx.shared import Inches, Pt

    opts = [_clean_mcq_option(o, i) for i, o in enumerate((options or [])[:4])]
    if not opts:
        return
    tbl = cell.add_table(rows=1, cols=len(opts))
    tbl.style = "Table Grid"
    _set_col_widths(tbl, [1.45] * len(opts))
    for j, opt in enumerate(opts):
        c = tbl.rows[0].cells[j]
        c.text = ""
        _para(c, f"({chr(97 + j)}) {opt}", size=11, space_after=0)


def _build_board_docx(doc, paper, meta):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    section_meta = paper.get("sectionMeta") or {}
    short_meta = section_meta.get("short") or {}
    long_meta = section_meta.get("long") or {}
    long_start = int(long_meta.get("longStartNo") or 2 + len(short_meta.get("blocks") or []) or 5)

    institute = (meta.get("institute") or "MUKABBIR COLLEGE GUJRAT").upper()
    _para(doc, institute, bold=True, size=17, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    _para(
        doc,
        meta.get("examTitle") or paper.get("title") or "Examination Paper",
        bold=True,
        size=13,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=8,
    )

    class_name = meta.get("className") or "—"
    year = "2nd year" if re.search(r"12|2nd|second", class_name, re.I) else "1st year"
    program = "Intermediate Part 2" if year == "2nd year" else "Intermediate Part 1"

    meta_table = doc.add_table(rows=3, cols=2)
    meta_table.style = "Table Grid"
    _set_col_widths(meta_table, [3.25, 3.25])
    meta_rows = [
        (("Program", program), ("Class", year)),
        (("Subject", meta.get("subject") or "—"), ("Time Allowed", meta.get("duration") or paper.get("duration") or "—")),
        (("Teacher", meta.get("teacherName") or "................................"), ("Maximum Marks", str(paper.get("marks") or "—"))),
    ]
    for i, ((l1, v1), (l2, v2)) in enumerate(meta_rows):
        row = meta_table.rows[i]
        row.cells[0].text = ""
        _para(row.cells[0], f"{l1}\n{v1}", size=11, space_after=0)
        row.cells[1].text = ""
        _para(row.cells[1], f"{l2}\n{v2}", size=11, space_after=0)

    doc.add_paragraph()
    name_p = doc.add_paragraph()
    name_p.paragraph_format.space_after = Pt(10)
    r1 = name_p.add_run("Name: ....................................................................")
    r1.font.name = FONT_NAME
    r1.font.size = Pt(11)
    r2 = name_p.add_run("          Roll No: ............................")
    r2.font.name = FONT_NAME
    r2.font.size = Pt(11)

    for sec in paper.get("sections") or []:
        key = _section_key(sec)
        qs = sec.get("questions") or []
        if not qs:
            continue

        doc.add_paragraph()

        if key == "objective":
            obj = section_meta.get("objective") or {}
            head = obj.get("boardHead") or "Q.No.1: Encircle the Correct Option."
            marks = obj.get("marksLine") or f"(1×{len(qs)}={len(qs)})"
            table = doc.add_table(rows=1 + len(qs), cols=2)
            table.style = "Table Grid"
            _set_col_widths(table, [0.45, 6.05])
            hdr = _merge_row(table, 0)
            _shade_cell(hdr)
            _para(hdr, f"{head}  {marks}", bold=True, size=11, space_after=0)
            for i, q in enumerate(qs, start=1):
                row = table.rows[i]
                row.cells[0].text = ""
                _para(row.cells[0], str(i), bold=True, size=11, space_after=0)
                cell = row.cells[1]
                cell.text = ""
                _para(cell, q.get("questionText") or "", size=11, space_after=3)
                _add_mcq_options_cell(cell, q.get("options"))

        elif key == "short":
            blocks = short_meta.get("blocks") or [
                {
                    "questionNo": 2,
                    "boardHead": "Q.No.2: Write Short Answers to any SIX questions.",
                    "printed": len(qs),
                    "marksLine": "",
                }
            ]
            _para(doc, "(Subjective Section)", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
            offset = 0
            groups = []
            for blk in blocks:
                slice_q = qs[offset : offset + int(blk.get("printed") or 0)]
                offset += len(slice_q)
                if slice_q:
                    groups.append((blk, slice_q))
            total_rows = len(groups) + sum(len(s) for _, s in groups)
            table = doc.add_table(rows=max(total_rows, 1), cols=2)
            table.style = "Table Grid"
            _set_col_widths(table, [0.45, 6.05])
            row_idx = 0
            for blk, slice_q in groups:
                hdr = table.rows[row_idx].cells[0]
                hdr.merge(table.rows[row_idx].cells[1])
                _shade_cell(hdr)
                head = blk.get("boardHead") or f"Q.No.{blk.get('questionNo', 2)}:"
                marks = blk.get("marksLine") or ""
                _para(hdr, f"{head}  {marks}".strip(), bold=True, size=11, space_after=0)
                row_idx += 1
                for j, q in enumerate(slice_q):
                    row = table.rows[row_idx]
                    row.cells[0].text = ""
                    roman = ROMAN[j] if j < len(ROMAN) else str(j + 1)
                    _para(row.cells[0], f"{roman}.", bold=True, size=11, space_after=0)
                    row.cells[1].text = ""
                    _para(row.cells[1], q.get("questionText") or "", size=11, space_after=0)
                    row_idx += 1

        elif key == "long":
            head = long_meta.get("boardHead") or "SECTION — II (Long Questions)"
            inst = long_meta.get("instruction") or ""
            marks = long_meta.get("marksLine") or ""
            table = doc.add_table(rows=1 + len(qs), cols=2)
            table.style = "Table Grid"
            _set_col_widths(table, [0.45, 6.05])
            hdr = _merge_row(table, 0)
            _shade_cell(hdr)
            _para(hdr, f"{head}  {inst}  {marks}".strip(), bold=True, size=11, space_after=0)
            for i, q in enumerate(qs):
                num = long_start + i
                row = table.rows[i + 1]
                row.cells[0].text = ""
                _para(row.cells[0], str(num), bold=True, size=11, space_after=0)
                row.cells[1].text = ""
                _para(row.cells[1], f"Q.No.{num}: {q.get('questionText') or ''}", size=11, space_after=0)

    return doc


def build_paper_docx(paper: dict, meta: dict | None = None) -> bytes:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ImportError as exc:
        raise RuntimeError("Install python-docx: pip install python-docx") from exc

    from io import BytesIO

    meta = meta or {}
    doc = Document()
    _setup_document(doc)

    if meta.get("boardFormat"):
        _build_board_docx(doc, paper, meta)
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    title = meta.get("examTitle") or paper.get("title") or "Practice Paper"
    _para(doc, title, bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    simple = meta.get("simpleLayout")
    if not simple:
        _para(
            doc,
            meta.get("institute") or "MUKABBIR COLLEGE GUJRAT",
            size=11,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=8,
        )
        table = doc.add_table(rows=4, cols=4)
        table.style = "Table Grid"
        rows_data = [
            ("Subject", meta.get("subject") or "—", "Date", meta.get("date") or "—"),
            ("Student Name", meta.get("studentName") or "________________", "Roll No", meta.get("rollNo") or "________________"),
            ("Time Allowed", meta.get("duration") or paper.get("duration") or "—", "Total Marks", str(paper.get("marks") or "—")),
            ("Class", meta.get("className") or "—", "Section", meta.get("section") or "—"),
        ]
        for row_idx, (l1, v1, l2, v2) in enumerate(rows_data):
            row = table.rows[row_idx]
            row.cells[0].text = str(l1)
            row.cells[1].text = str(v1)
            row.cells[2].text = str(l2)
            row.cells[3].text = str(v2)

    section_meta = paper.get("sectionMeta") or {}
    for sec in paper.get("sections") or []:
        doc.add_paragraph()
        key = _section_key(sec)
        sm = section_meta.get(key) or {}
        label = sm.get("title") or sec.get("title") or "Section"
        _para(doc, label.upper(), bold=True, size=12, space_after=4)

        qs = sec.get("questions") or []
        if key == "short" and sm.get("blocks"):
            offset = 0
            for blk in sm["blocks"]:
                slice_q = qs[offset : offset + int(blk.get("printed") or 0)]
                offset += len(slice_q)
                _para(doc, f"{blk.get('boardHead') or ''} {blk.get('marksLine') or ''}".strip(), bold=True, size=11)
                for j, q in enumerate(slice_q):
                    roman = ROMAN[j] if j < len(ROMAN) else str(j + 1)
                    _para(doc, f"{roman}. {q.get('questionText') or ''}", size=11)
        else:
            start = int((section_meta.get("long") or {}).get("longStartNo") or 5)
            for q_idx, q in enumerate(qs, start=1 if key != "long" else start):
                num = q_idx if key != "long" else start + (q_idx - start)
                _para(doc, f"{num}. {q.get('questionText') or ''}", size=11)
                if (q.get("questionType") or "").lower() == "mcq":
                    opts = q.get("options") or []
                    line = "    ".join(f"({chr(97 + i)}) {_clean_mcq_option(o, i)}" for i, o in enumerate(opts[:4]))
                    if line:
                        _para(doc, line, size=10)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
